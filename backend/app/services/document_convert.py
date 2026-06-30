from __future__ import annotations

import io
import os
import re
import tempfile
import zipfile
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt

HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
BULLET_PATTERN = re.compile(r"^[-*+]\s+(.+)$")
ORDERED_PATTERN = re.compile(r"^\d+\.\s+(.+)$")
BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
ITALIC_PATTERN = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


def _add_formatted_runs(paragraph, text: str) -> None:
    if not text:
        return

    parts: list[tuple[str, bool, bool]] = []
    remaining = text
    while remaining:
        bold_match = BOLD_PATTERN.search(remaining)
        italic_match = ITALIC_PATTERN.search(remaining)
        match = None
        if bold_match and (not italic_match or bold_match.start() <= italic_match.start()):
            match = ("bold", bold_match)
        elif italic_match:
            match = ("italic", italic_match)

        if not match:
            parts.append((remaining, False, False))
            break

        kind, found = match
        if found.start() > 0:
            parts.append((remaining[: found.start()], False, False))
        parts.append((found.group(1), kind == "bold", kind == "italic"))
        remaining = remaining[found.end() :]

    for segment, bold, italic in parts:
        run = paragraph.add_run(segment)
        run.bold = bold
        run.italic = italic


def markdown_to_docx(content: str, title: str = "문서") -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        heading = HEADING_PATTERN.match(line)
        if heading:
            level = min(len(heading.group(1)), 6)
            doc.add_heading(heading.group(2).strip(), level=level)
            continue

        bullet = BULLET_PATTERN.match(line)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(paragraph, bullet.group(1).strip())
            continue

        ordered = ORDERED_PATTERN.match(line)
        if ordered:
            paragraph = doc.add_paragraph(style="List Number")
            _add_formatted_runs(paragraph, ordered.group(1).strip())
            continue

        paragraph = doc.add_paragraph()
        _add_formatted_runs(paragraph, line.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _hwpx_section_xml(paragraphs: list[str]) -> str:
    para_xml = []
    for text in paragraphs:
        safe = escape(text)
        para_xml.append(
            f'<hp:p><hp:run><hp:t>{safe}</hp:t></hp:run></hp:p>'
        )
    body = "".join(para_xml) or '<hp:p><hp:run><hp:t></hp:t></hp:run></hp:p>'
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" '
        'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">'
        f"{body}</hs:sec>"
    )


def markdown_to_hwpx(content: str, title: str = "문서") -> bytes:
    paragraphs: list[str] = [title]
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = HEADING_PATTERN.match(line)
        if heading:
            paragraphs.append(heading.group(2).strip())
            continue
        bullet = BULLET_PATTERN.match(line)
        if bullet:
            paragraphs.append(f"• {bullet.group(1).strip()}")
            continue
        ordered = ORDERED_PATTERN.match(line)
        if ordered:
            paragraphs.append(ordered.group(1).strip())
            continue
        paragraphs.append(line)

    section_xml = _hwpx_section_xml(paragraphs)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("mimetype", "application/hwp+zip")
        archive.writestr(
            "version.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version" '
            'tagetApplication="WORD_PROCESSOR" major="5" minor="1" micro="0" buildNumber="0"/>',
        )
        archive.writestr(
            "Contents/header.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" '
            'version="1.5" secCnt="1"><hh:beginNum page="1" footnote="1" endnote="1" pic="1" tbl="1" equation="1"/>'
            '<hh:refList><hh:fontfaces itemCnt="1"><hh:fontface lang="HANGUL" fontCnt="1">'
            '<hh:font id="0" face="맑은 고딕" type="TTF" isEmbedded="0"/></hh:fontface></hh:fontfaces></hh:refList></hh:head>',
        )
        archive.writestr("Contents/section0.xml", section_xml)
        archive.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container" '
            'ocf:version="1.0"><ocf:rootfiles><ocf:rootfile full-path="Contents/header.xml" '
            'media-type="application/hwpml+xml"/><ocf:rootfile full-path="Contents/section0.xml" '
            'media-type="application/hwpml+xml"/></ocf:rootfiles></ocf:container>',
        )
        archive.writestr(
            "META-INF/manifest.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<odf:manifest xmlns:odf="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" '
            'odf:version="1.2"><odf:file-entry odf:full-path="/" odf:media-type="application/hwp+zip"/>'
            '<odf:file-entry odf:full-path="Contents/header.xml" odf:media-type="application/hwpml+xml"/>'
            '<odf:file-entry odf:full-path="Contents/section0.xml" odf:media-type="application/hwpml+xml"/>'
            '</odf:manifest>',
        )
        preview = "\n".join(paragraphs[:20])
        archive.writestr("Preview/PrvText.txt", preview.encode("utf-8"))

    return buffer.getvalue()


def plain_text_to_docx(content: str, title: str = "문서") -> bytes:
    return markdown_to_docx(content, title=title)


MIME_BY_FORMAT = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "hwpx": "application/hwp+zip",
    "hwp": "application/x-hwp",
}


def convert_content(content: str, target_format: str, title: str = "문서") -> tuple[bytes, str]:
    normalized = target_format.lower().strip()
    if normalized in {"docx", "word"}:
        return markdown_to_docx(content, title=title), "docx"
    if normalized in {"hwp", "hwpx", "한글"}:
        return markdown_to_hwpx(content, title=title), "hwpx"
    raise ValueError(f"지원하지 않는 형식입니다: {target_format}. docx 또는 hwpx(한글)를 사용하세요.")


TEXT_FORMATS = frozenset({"markdown", "md", "txt", "text"})
BINARY_TEXT_FORMATS = frozenset({"docx", "word", "hwpx", "hwp"})


def _decode_text_bytes(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="replace")


def _is_zip_bytes(data: bytes) -> bool:
    return len(data) >= 2 and data[:2] == b"PK"


def _is_ole_bytes(data: bytes) -> bool:
    return len(data) >= 8 and data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def extract_hwp_text(data: bytes) -> str:
    from hwpkit import extract_text_from_hwp

    with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        text = extract_text_from_hwp(tmp_path).strip()
        if not text:
            raise ValueError("한글 문서에서 추출할 수 있는 텍스트가 없습니다.")
        return text
    except Exception as exc:
        message = str(exc)
        if "OLE" in type(exc).__name__ or "ole" in message.lower():
            raise ValueError("올바른 한글(.hwp) 문서 파일이 아닙니다.") from exc
        raise ValueError(f"한글 문서를 읽을 수 없습니다: {message}") from exc
    finally:
        os.unlink(tmp_path)


def extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def extract_hwpx_text(data: bytes) -> str:
    if _is_ole_bytes(data):
        return extract_hwp_text(data)
    if not _is_zip_bytes(data):
        raise ValueError("올바른 한글(.hwpx) 문서 파일이 아닙니다.")

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            preview_path = "Preview/PrvText.txt"
            if preview_path in archive.namelist():
                return _decode_text_bytes(archive.read(preview_path)).strip()

            section_path = "Contents/section0.xml"
            if section_path not in archive.namelist():
                return ""

            xml = archive.read(section_path).decode("utf-8", errors="replace")
            texts = re.findall(r"<hp:t>(.*?)</hp:t>", xml)
            unescaped = [t.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&") for t in texts]
            return "\n".join(line for line in unescaped if line.strip())
    except zipfile.BadZipFile as exc:
        raise ValueError("한글 문서 파일을 읽을 수 없습니다. 파일이 손상되었을 수 있습니다.") from exc


def extract_text_from_bytes(data: bytes, file_format: str) -> str:
    normalized = file_format.lower().strip()
    if normalized in TEXT_FORMATS:
        return _decode_text_bytes(data)
    if normalized in {"docx", "word"}:
        if not _is_zip_bytes(data):
            raise ValueError("올바른 Word(.docx) 문서 파일이 아닙니다.")
        return extract_docx_text(data)
    if normalized == "hwp":
        return extract_hwp_text(data)
    if normalized == "hwpx":
        return extract_hwpx_text(data)
    return _decode_text_bytes(data)
